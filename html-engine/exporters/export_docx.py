"""DOM-walking HTML -> DOCX exporter.

Opens a document.html in headless chromium, walks the `.doc-page`
block flow in document order, and emits a python-docx file where each
block becomes a real Word construct: headings, paragraphs, bullet and
numbered lists, block quotes, shaded callouts, tables, and horizontal
rules. Inline <strong>/<em>/<code>/<a> become formatted runs.

Usage:
    python3 export_docx.py --url http://host/preview/.../doc-v1.html \\
                            --out  /path/to/doc.docx

The document side of Oneshot is deliberately block-flow only (no
floats, multi-column, flex/grid, or absolute positioning), which is
exactly the subset Word represents natively. That discipline is what
makes this export one-for-one instead of a rasterized approximation.
The page's own CSS provides the accent color and the class vocabulary
(.eyebrow / .lede / .bullets / .callout|.alert / .doc-meta / etc.);
unknown blocks degrade gracefully to plain paragraphs.
"""
from __future__ import annotations
import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# JS that runs inside the rendered page. Returns:
#   { accent, ink, muted, blocks: [ {type, ...}, ... ] }
# Each block carries `runs` (inline run segments) where applicable.
DOM_WALK_JS = r"""
(() => {
  const page = document.querySelector('.doc-page') || document.body;
  const cs = getComputedStyle(document.documentElement);
  const accent = (cs.getPropertyValue('--accent') || '#0071E3').trim();
  const ink    = (cs.getPropertyValue('--ink') || '#1D1D1F').trim();
  const muted  = (cs.getPropertyValue('--stone') || '#6E6E73').trim();

  const hasClass = (el, c) => el.classList && el.classList.contains(c);

  // Flatten an element's descendants into formatted run segments,
  // honoring strong/b/em/i/code/a along the way.
  function runs(el) {
    const out = [];
    function walk(node, fmt) {
      for (const child of node.childNodes) {
        if (child.nodeType === 3) {            // text node
          const t = child.textContent;
          if (t && t.trim() !== '' || t === ' ') out.push({ text: t, ...fmt });
        } else if (child.nodeType === 1) {     // element
          const tag = child.tagName.toLowerCase();
          const next = { ...fmt };
          if (tag === 'strong' || tag === 'b' || hasClass(child, 'em')) next.bold = true;
          if (tag === 'em' || tag === 'i') next.italic = true;
          if (tag === 'code' || tag === 'kbd' || hasClass(child, 'mono')) next.mono = true;
          if (tag === 'a') next.link = child.getAttribute('href') || '';
          if (tag === 'br') { out.push({ text: '\n', ...fmt }); continue; }
          walk(child, next);
        }
      }
    }
    walk(el, {});
    // Collapse runs of pure-whitespace to keep Word tidy.
    return out
      .map(r => ({ ...r, text: r.text.replace(/\s+/g, ' ') }))
      .filter((r, i, a) => !(r.text === ' ' && (i === 0 || i === a.length - 1)));
  }

  const text = (el) => (el.textContent || '').replace(/\s+/g, ' ').trim();

  function listItems(el) {
    return Array.from(el.querySelectorAll(':scope > li')).map(li => runs(li));
  }

  function tableData(el) {
    const rows = Array.from(el.querySelectorAll('tr'));
    let headers = [];
    const thead = el.querySelector('thead tr') || (rows[0] && rows[0].querySelector('th') ? rows[0] : null);
    if (thead) headers = Array.from(thead.querySelectorAll('th,td')).map(text);
    const bodyRows = rows
      .filter(r => r !== thead)
      .map(r => Array.from(r.querySelectorAll('td,th')).map(text))
      .filter(r => r.length);
    return { headers, rows: bodyRows };
  }

  const blocks = [];
  for (const el of Array.from(page.children)) {
    const tag = el.tagName.toLowerCase();
    if (el.offsetParent === null && tag !== 'hr' && getComputedStyle(el).display === 'none') continue;

    if (tag === 'h1' || hasClass(el, 'title')) { blocks.push({ type: 'h1', runs: runs(el) }); continue; }
    if (tag === 'h2') { blocks.push({ type: 'h2', runs: runs(el) }); continue; }
    if (tag === 'h3') { blocks.push({ type: 'h3', runs: runs(el) }); continue; }
    if (hasClass(el, 'eyebrow') || hasClass(el, 'kicker')) { blocks.push({ type: 'eyebrow', text: text(el) }); continue; }
    if (hasClass(el, 'lede')) { blocks.push({ type: 'lede', runs: runs(el) }); continue; }
    if (hasClass(el, 'doc-meta')) {
      blocks.push({ type: 'meta', items: Array.from(el.children).map(text).filter(Boolean) || [text(el)] });
      continue;
    }
    if (tag === 'hr') { blocks.push({ type: 'hr' }); continue; }
    if (tag === 'blockquote' || hasClass(el, 'quote')) { blocks.push({ type: 'quote', runs: runs(el) }); continue; }
    if (hasClass(el, 'callout') || hasClass(el, 'alert')) {
      let variant = 'info';
      if (hasClass(el, 'warn')) variant = 'warn';
      else if (hasClass(el, 'negative')) variant = 'negative';
      blocks.push({ type: 'callout', variant, runs: runs(el) });
      continue;
    }
    if (tag === 'ul' || hasClass(el, 'bullets')) { blocks.push({ type: 'list', ordered: false, items: listItems(el) }); continue; }
    if (tag === 'ol') { blocks.push({ type: 'list', ordered: true, items: listItems(el) }); continue; }
    if (tag === 'table' || hasClass(el, 'doc-table')) { blocks.push({ type: 'table', ...tableData(el) }); continue; }
    if (tag === 'pre') { blocks.push({ type: 'pre', text: el.textContent || '' }); continue; }
    // Unknown wrappers (section/div/article): recurse one level so a
    // grouping container doesn't swallow its children.
    if ((tag === 'div' || tag === 'section' || tag === 'article') && el.children.length) {
      for (const inner of Array.from(el.children)) {
        const itag = inner.tagName.toLowerCase();
        if (itag === 'h2') blocks.push({ type: 'h2', runs: runs(inner) });
        else if (itag === 'h3') blocks.push({ type: 'h3', runs: runs(inner) });
        else if (itag === 'ul') blocks.push({ type: 'list', ordered: false, items: listItems(inner) });
        else if (itag === 'ol') blocks.push({ type: 'list', ordered: true, items: listItems(inner) });
        else if (itag === 'table') blocks.push({ type: 'table', ...tableData(inner) });
        else if (text(inner)) blocks.push({ type: 'p', runs: runs(inner) });
      }
      continue;
    }
    if (text(el)) { blocks.push({ type: 'p', runs: runs(el) }); }
  }

  return { accent, ink, muted, blocks };
})()
"""

SANS = "Helvetica Neue"
MONO = "Consolas"


def hex_to_rgb(h: str) -> RGBColor:
    h = (h or "").strip().lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    if len(h) != 6:
        return RGBColor(0x1D, 0x1D, 0x1F)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_shading(cell, hex_fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill.lstrip("#"))
    tcPr.append(shd)


def add_runs(paragraph, runs, accent: RGBColor, base_size: float, base_color: RGBColor):
    """Append formatted run segments to a paragraph."""
    if not runs:
        return
    for seg in runs:
        text = seg.get("text", "")
        if text == "\n":
            paragraph.add_run().add_break()
            continue
        r = paragraph.add_run(text)
        r.font.size = Pt(base_size)
        r.font.name = MONO if seg.get("mono") else SANS
        r.font.color.rgb = base_color
        if seg.get("bold"):
            r.bold = True
        if seg.get("italic"):
            r.italic = True
        if seg.get("mono"):
            r.font.size = Pt(base_size - 1)
        if seg.get("link"):
            r.font.color.rgb = accent
            r.font.underline = True


# CSS px (96dpi) -> Word points. Documents are sized for reading at
# literal pixel size, so the px values map straight to print points.
PX_TO_PT = {
    "h1": 26, "h2": 18, "h3": 14, "lede": 14, "body": 11,
    "eyebrow": 9, "small": 9.5, "quote": 13,
}

CALLOUT_FILL = {"info": "EBF3FD", "warn": "FFF8EC", "negative": "FCEDEB"}
CALLOUT_BORDER = {"info": "2997FF", "warn": "C97F00", "negative": "D93025"}


def build(data: dict, out_path: str):
    accent = hex_to_rgb(data.get("accent", "#0071E3"))
    ink = hex_to_rgb(data.get("ink", "#1D1D1F"))
    muted = hex_to_rgb(data.get("muted", "#6E6E73"))

    doc = Document()
    # Letter page with a comfortable reading margin.
    sec = doc.sections[0]
    sec.left_margin = Inches(1.1)
    sec.right_margin = Inches(1.1)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = SANS
    normal.font.size = Pt(PX_TO_PT["body"])
    normal.font.color.rgb = ink
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.3

    for b in data.get("blocks", []):
        t = b.get("type")

        if t == "eyebrow":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run((b.get("text") or "").upper())
            r.font.name = SANS
            r.font.size = Pt(PX_TO_PT["eyebrow"])
            r.bold = True
            r.font.color.rgb = accent
            _set_tracking(r, 120)

        elif t == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.08
            add_runs(p, b.get("runs"), accent, PX_TO_PT["h1"], ink)
            for r in p.runs:
                r.bold = True

        elif t == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            add_runs(p, b.get("runs"), accent, PX_TO_PT["h2"], ink)
            for r in p.runs:
                r.bold = True

        elif t == "h3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, b.get("runs"), accent, PX_TO_PT["h3"], ink)
            for r in p.runs:
                r.bold = True

        elif t == "lede":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            add_runs(p, b.get("runs"), accent, PX_TO_PT["lede"], muted)

        elif t == "meta":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            items = b.get("items") or []
            for i, it in enumerate(items):
                if i:
                    sep = p.add_run("    •    ")
                    sep.font.size = Pt(PX_TO_PT["small"])
                    sep.font.color.rgb = muted
                r = p.add_run(it)
                r.font.size = Pt(PX_TO_PT["small"])
                r.font.color.rgb = muted
            _bottom_border(p)

        elif t == "p":
            p = doc.add_paragraph()
            add_runs(p, b.get("runs"), accent, PX_TO_PT["body"], ink)

        elif t == "list":
            style = "List Number" if b.get("ordered") else "List Bullet"
            for item in b.get("items") or []:
                p = doc.add_paragraph(style=style)
                p.paragraph_format.space_after = Pt(3)
                add_runs(p, item, accent, PX_TO_PT["body"], ink)

        elif t == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            add_runs(p, b.get("runs"), accent, PX_TO_PT["quote"], muted)
            for r in p.runs:
                r.italic = True
            _left_bar(p, data.get("accent", "#0071E3"))

        elif t == "callout":
            variant = b.get("variant", "info")
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell = tbl.rows[0].cells[0]
            set_cell_shading(cell, CALLOUT_FILL.get(variant, "EBF3FD"))
            _cell_borders(cell, CALLOUT_BORDER.get(variant, "2997FF"))
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            add_runs(cp, b.get("runs"), accent, PX_TO_PT["body"], ink)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif t == "table":
            headers = b.get("headers") or []
            rows = b.get("rows") or []
            ncol = max([len(headers)] + [len(r) for r in rows] or [1]) or 1
            tbl = doc.add_table(rows=0, cols=ncol)
            tbl.style = "Light Grid Accent 1"
            if headers:
                cells = tbl.add_row().cells
                for i in range(ncol):
                    para = cells[i].paragraphs[0]
                    r = para.add_run(headers[i] if i < len(headers) else "")
                    r.bold = True
                    r.font.size = Pt(PX_TO_PT["small"])
            for row in rows:
                cells = tbl.add_row().cells
                for i in range(ncol):
                    para = cells[i].paragraphs[0]
                    r = para.add_run(row[i] if i < len(row) else "")
                    r.font.size = Pt(PX_TO_PT["small"])
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif t == "pre":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(b.get("text") or "")
            r.font.name = MONO
            r.font.size = Pt(PX_TO_PT["small"])

        elif t == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            _bottom_border(p)

    doc.save(out_path)


def _set_tracking(run, twentieths):
    """Letter-spacing in twentieths of a point (Word's w:spacing)."""
    rPr = run._element.get_or_add_rPr()
    el = OxmlElement("w:spacing")
    el.set(qn("w:val"), str(twentieths))
    rPr.append(el)


def _bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "D2D2D7")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _left_bar(paragraph, accent_hex):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), accent_hex.lstrip("#"))
    pbdr.append(left)
    pPr.append(pbdr)


def _cell_borders(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color_hex.lstrip("#"))
        borders.append(e)
    tcPr.append(borders)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--url", required=True, help="rendered document preview URL")
    ap.add_argument("--out", required=True, help="output .docx path")
    ap.add_argument("--timeout", type=int, default=30000)
    args = ap.parse_args()

    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(args=["--no-sandbox"])
        page = browser.new_page(viewport={"width": 820, "height": 1200})
        page.goto(args.url, wait_until="networkidle", timeout=args.timeout)
        page.wait_for_timeout(350)
        data = page.evaluate(DOM_WALK_JS)
        browser.close()

    if not data or not data.get("blocks"):
        print("no .doc-page content found", file=sys.stderr)
        sys.exit(2)

    build(data, args.out)
    print(json.dumps({"ok": True, "blocks": len(data["blocks"]), "out": args.out}))


if __name__ == "__main__":
    main()
